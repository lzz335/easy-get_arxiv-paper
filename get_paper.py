import urllib.request
import feedparser
import numpy as np
import pandas as pd
from docx import Document
import docx.shared
from datetime import datetime
import os

# 基本参数的设置
base_url = 'http://export.arxiv.org/api/query?'
start = 0
max_results = 2000
folder_name = "output"  # 输出的文件夹，如果不存在会自动生成

# 需要的检索结果配置，注释掉自己不需要的条目
need_res = [
    'arxiv-id',
    'Published',
    'Title',
    'Last Author',
    'Authors',
    'abs page link',
    'pdf link',
    'Journal reference',
    'Comments',
    'Primary Category',
    'All Categories',
    'Abstract',
]

'''
|prefix| explanation |
| ---- |-------------|
|ti| 标题          |
|au| 作者          |
|abs| 摘要          |
|co| 会议          |
|jr|Journal Reference|
|cat|Subject Category|
|rn|Report Number|
|id|Id (use id_list instead)|
|all|All of the above|
'''

# 关于搜索的细节
co_filter = True  # 决定是否启用会议搜索的filter，用于解决在或关系中限定的会议但是没有效果的问题
# 关于上面的问题说明：latex提供的公共接口在以或关系为条件进行检索的时候可能会出现引入其他杂七杂八的会议的情况，本开关一般推荐启用

Main_search = 'abs:Embedding'
And_search = [
    'abs:Search',
    'abs:System'
]

Or_search = [
    'co:ACL',
    'co:EMNLP',
    'co:NAACL',
    'co:NeurIPS',
    'co:ICML',
    'co:NIPS',
    'co:COLING'
]
ANDNOT_search = [
]
save_name = 'Embedding 2023'

search_query = Main_search.replace(' ', '+')
# 备注，在请求的时候需要使用'+'来替代' '
for text in And_search:
    search_query = search_query + '+AND+' + text.replace(' ', '+')

for text in Or_search:
    search_query = search_query + '+OR+' + text.replace(' ', '+')

for text in ANDNOT_search:
    search_query = search_query + '+ANDNOT+' + text.replace(' ', '+')

query = 'search_query=%s&start=%i&max_results=%i' % (search_query,
                                                     start,
                                                     max_results)

# 发送请求并获得响应
response = urllib.request.urlopen(base_url + query).read()
feed = feedparser.parse(response)

# 本部分用于初始化filter
filter_dict = []
if co_filter:  # 取决于前面设置是否启用filter
    for item in Or_search:
        if 'co:' in item:
            filter_dict.append(item.split(':')[1])  # 将检索的会议名称加入字典列表


def is_in_coList(co_dict: list, co_name: str):
    for co in co_dict:
        if co in co_name:
            return True
    return False


# 输出的excel表头
search_res = [need_res]

# 生成word文件
doc = Document()


def add_information(title_list, search_item, res_item, output_list):
    if search_item in title_list:
        output_list.append(res_item.replace('\n', ''))
    return output_list


def setFirstPage(document: Document(), message):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(message)
    run.bold = True
    run.font.size = docx.shared.Pt(14)
    paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER


# 第一页的标识
setFirstPage(doc, "检索文件名称" + save_name)
current_time = datetime.now()
formatted_time = current_time.strftime("%Y-%m-%d %H:%M:%S")
setFirstPage(doc, "检索时间：" + formatted_time)
setFirstPage(doc, "检索逻辑：" + query)

# 本部分的程序用于处理在输入或关系的时候，如果限定词为co，有可能会搜索出一些不在列表里的paper


for entry in feed.entries:

    # 检查会议是否满足条件
    try:
        comment = entry.arxiv_comment
    except AttributeError:
        comment = 'No comment found'

    if len(filter_dict) != 0 and co_filter:
        # 当存在检索的会议且启动filter时
        if not is_in_coList(co_dict=filter_dict, co_name=comment):
            continue

    doc.add_page_break()

    # 添加标题（居中，字体大小大于摘要）
    setFirstPage(doc, entry.title.replace('\n', ' '))

    one_result = []
    add_information(title_list=need_res, search_item='arxiv-id', res_item=entry.id.split('/abs/')[-1],
                    output_list=one_result)
    add_information(title_list=need_res, search_item='Published', res_item=entry.published, output_list=one_result)
    add_information(title_list=need_res, search_item='Title', res_item=entry.title, output_list=one_result)

    author_string = entry.author
    try:
        author_string += ' (%s)' % entry.arxiv_affiliation
    except AttributeError:
        pass

    add_information(title_list=need_res, search_item='Last Author', res_item=author_string, output_list=one_result)

    author_names = ''
    try:
        author_names = ', '.join(author.name for author in entry.authors)
    except AttributeError:
        pass

    add_information(title_list=need_res, search_item='Authors', res_item=author_names, output_list=one_result)

    abs_link = ''
    pdf_link = ''
    for link in entry.links:
        if link.rel == 'alternate':
            abs_link = link.href

        elif link.title == 'pdf':
            pdf_link = link.href

    add_information(title_list=need_res, search_item='abs page link', res_item=abs_link, output_list=one_result)
    add_information(title_list=need_res, search_item='pdf link', res_item=pdf_link, output_list=one_result)
    setFirstPage(doc, "arxiv摘要页地址：" + abs_link)
    try:
        journal_ref = entry.arxiv_journal_ref
    except AttributeError:
        journal_ref = 'No journal ref found'
    add_information(title_list=need_res, search_item='Journal reference', res_item=journal_ref, output_list=one_result)

    add_information(title_list=need_res, search_item='Comments', res_item=comment, output_list=one_result)
    setFirstPage(doc, "会议名字：" + comment)

    add_information(title_list=need_res, search_item='Primary Category', res_item=entry.tags[0]['term'],
                    output_list=one_result)

    all_categories = [t['term'] for t in entry.tags]
    add_information(title_list=need_res, search_item='All Categories', res_item=', '.join(all_categories),
                    output_list=one_result)
    add_information(title_list=need_res, search_item='Abstract', res_item=entry.summary, output_list=one_result)

    # 添加摘要（换行，字体大小较小）
    abstract_paragraph = doc.add_paragraph(entry.summary.replace('\n', ' '))
    abstract_paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT
    abstract_run = abstract_paragraph.runs[0]
    abstract_run.font.size = docx.shared.Pt(12)

    search_res.append(one_result)

res = np.array(search_res)

df = pd.DataFrame(res)

for paragraph in doc.paragraphs:
    paragraph.paragraph_format.line_spacing = 1
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'

if not os.path.exists(folder_name):
    os.makedirs(folder_name)
    print(f"文件夹 '{folder_name}' 创建成功！")

df.to_excel(folder_name+'/' + save_name + '.xlsx', index=False, header=False)
print('输出成功 %s.xlsx' % save_name)
doc.save(folder_name+'/' + save_name + ".docx")
print('输出成功 %s.docx' % save_name)
